#!/usr/bin/env python3

'''
MOU 파일 생성

USAGE:
  $> python3 doc_mou.py
'''

# docx 모듈 로드
from docx import Document
# 정렬 상수
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt

def make_mou(filename, data, templete):
    """MOU 문서 생성.

    기존의 파일을 이용해서 새로운 문서를 생성한다.
    기존 문서에서 교체할 부분은 `{{key}}'형태로 
    지정이 되어 있다.

    """
    # 문서 생성
    document = Document(templete)

    # 모든 문단을 찾아서
    for p in document.paragraphs:
        for i in p.runs:
            # 각 run의 text값을 바꾼다.
            s = i.text
            for k in data:
                s = s.replace('{{%s}}' % k, str(data[k]))
            i.text = s

    # 모든 테이블을 찾아서
    for t in document.tables:
        # 모든 cell에 있는 
        for r in t.rows:
            for c in r.cells:
                # 문단을 찾아서 
                for i in c.paragraphs:
                    for j in i.runs:
                        # 각 run의 text값을 바꾼다.
                        s = j.text
                        for k in data:
                            s = s.replace('{{%s}}' % k, str(data[k]))
                        j.text = s


    # 문서 저장
    document.save(filename)


def main():

    # MOU 문서 생성에 필요한 데이터
    data = {
        "company1_name": "회사1",
        "company1_ceo" : "홍길동",
        "company2_name": "회사2",
        "company2_ceo" : "허균",
        "year": 2016,
        "month": 1,
        "day": 1,
        "condition1": "1. 두 회사의 모든 제품 정보를 교환한다.",
        "condition2": "2. 두 회사의 연구 결과를 공유한다.",
        "condition3": ""
    }

    make_mou(filename="mou.docx", data = data, templete="templete/MOU.docx")

if __name__=='__main__':
    main()

