#!/usr/bin/env python3

'''
MOU 파일 생성

USAGE:
  $> python3 doc_mou.py
'''

# docx 모듈 로드
from docx import Document
# 정렬 상수
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt

def make_mou(filename, data):
    """MOU 문서 생성.

    MOU 문서를 생성한다. 이 때 필요한 데이터(data)와 
    파일이름을 받는다.

    주의:
    기존에 filename과 동일한 파일이 이미 존재를 하면 
    기존 파일이 생성되고 새로운 파일로 대체된다.
    """

    # 문서 생성
    document = Document()

    # 제목 추가
    p = document.add_heading('MOU(양해각서)', level=1)
    # 제목을 중앙정렬로
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    
    text = "(주) {company1_name}와  (주){company2_name} 의 업무제휴에 관한 양해각서".format(**data)
    p = document.add_paragraph()
    p.add_run(text).bold = True


    text = "㈜ (이하 ‘갑’이라 함)과 ㈜ (이하 ‘을’이라 함)은 당사자간의 " \
           "우호협력관계를 확인하고 상호신뢰를 바탕으로 제휴에 따른 각자의 " \
           "책임을 인식하며, 인터넷 사업 분야에 있어 교류와 협력이 당사자간의 "\
           "상호 이해증진에 기여할 것임을 확신하면서, 다음과 같이 합의한다."
    p = document.add_paragraph(text)


    text = "제1조 목 적"
    p = document.add_paragraph(text, style="Heading 2")
    text = "본 계약은 ‘갑’과 ‘을’의 업무상 상호 공동이익의 증진을 도모함에 그 목적이 있다."
    p = document.add_paragraph(text)

    text = "제2조 협력내용"
    p = document.add_paragraph(text, style="Heading 2")
    text = "본 양해각서에 의한 제휴협력관계의 내용은 다음과 같다."
    p = document.add_paragraph(text)
    for cond in data["conditions"]:
        text = "   " + cond
        p = document.add_paragraph(text)    

        
    text = "제3조 분쟁해결"
    p = document.add_paragraph(text, style="Heading 2")
    text = "본 양해각서의 해석이나 적용에 관한 분쟁은 당사자 간의 상호협력에 의하여 우호적으로 해결하며, 제3자의 개입을 허용하지 않는다."
    p = document.add_paragraph(text)

    
    text = "제4조 의무사항"
    p = document.add_paragraph(text, style="Heading 2")
    text = '1)"갑" 과 "을"은 제휴업무를 수행함에 있어서 선량한 관리자로서의 주의를 다하여야 한다.'
    p = document.add_paragraph(text)
    text = '2) "갑" 과 "을"은 본 계약의 내용을 신의에 따라 성실하게 이행한다.'
    p = document.add_paragraph(text)


    text = "제5조 기밀유지"
    p = document.add_paragraph(text, style="Heading 2")
    text = '"갑" 과 "을"은 업무제휴를 통하여 취득한 정보 등을 상대방의 동의 없이 상호간의 공동 사업 추진 외의 목적에 사용하거나 외부에 누출 및 누설하여서는 아니되며, 계약 종료 이후에도 같다.'
    p = document.add_paragraph(text)

    text = "제6조 제반 업무 연락"
    p = document.add_paragraph(text, style="Heading 2")
    text = "본 협정과 관련된 제반 업무 연락은 문서로 함을 원칙으로 한다. 이 경우 문서는 인편 및 우편 그리고 팩스로 전달함으로써 정당하게 이루어진 것으로 한다. "
    p = document.add_paragraph(text)

    text = "제7조 계약의 해지"
    p = document.add_paragraph(text, style="Heading 2") 
    text = "본 협약은 서명일로부터 구체적인 본 계약을 체결하기 전까지 유효하다."
    p = document.add_paragraph(text)

    text = "제8조 유효기간"
    p = document.add_paragraph(text, style="Heading 2")
    text = "본 협약은 서명일로부터 구체적인 본 계약을 체결하기 전까지 유효하다."
    p = document.add_paragraph(text)

    text = "제9조 법적구속력"
    p = document.add_paragraph(text, style="Heading 2")
    text = "본 업무협약은 양사의 상호 업무에 관한 협력사항을 열거한 것으로, 제5조를 제외하고는 법적 구속력을 갖지 않는다."
    p = document.add_paragraph(text)
    
    text = "제10조 기타사항"
    p = document.add_paragraph(text, style="Heading 2")
    text = "본 계약서상에 명시되지 않은 기타 사항은 별도 협의 하에 처리한다. 본 양해각서는 {year}년 {month}월 {day}일 2부를 작성, 날인되었으며, 이 모두는 동등한 효력을 지닌다.".format(**data)
    p = document.add_paragraph(text)

    # 공백 추가
    p = document.add_paragraph("\n\n\n")

    table = document.add_table(rows=2, cols=2)
    cells = table.rows[0].cells
    cells[0].text = '갑: ㈜ {company1_name}'.format(**data)
    cells[1].text = '갑: ㈜ {company2_name}'.format(**data)    

    cells = table.rows[1].cells
    cells[0].text = '대표이사  {company1_ceo}'.format(**data)
    cells[1].text = '대표이사  {company2_ceo}'.format(**data)

    update_style(document)

    # 문서 저장
    document.save(filename)

def update_style(document):
    """문서의 스타일 재 설정.

    문서의 기본 스타일 중 다음과 같은 스타일을 재정의한다.

    - Heading 1
    - Heading 2
    - Normal
    """
    h1 = document.styles["Heading 1"]
    if h1:
        h1.paragraph_format.space_before = Pt(50)
        h1.paragraph_format.space_after = Pt(50)
        h1.font.size = Pt(26)
        
    h2 = document.styles["Heading 2"]
    if h2:
        h2.paragraph_format.space_before = Pt(30)
        h2.paragraph_format.space_after = Pt(15)

    s = document.styles["Normal"]
    if s:
        s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def main():

    # MOU 문서 생성에 필요한 데이터
    data = {
        "company1_name": "회사1",
        "company1_ceo" : "홍길동",
        "company2_name": "회사2",
        "company2_ceo" : "허균",
        "year": 2016,
        "month": 1,
        "day": 1,
        "conditions": [
            "1. 두 회사의 모든 제품 정보를 교환한다.",
            "2. 두 회사의 연구 결과를 공유한다."
        ]
    }

    make_mou(filename="mou.docx", data = data)

if __name__=='__main__':
    main()

