import os
import os.path
from docx import Document


def find_docx(target_dir):
    '주어진 폴더 아래 있는 모든 docx 파일을 찾아 리턴한다.'

    ret = []
    for root, dirs, files in os.walk(target_dir):
        for f in files:

            # docx 확장자가 아니면 skip한다.
            if not f.endswith('.docx'):
                continue

            # 절대 경로를 구한다.
            cur_file = os.path.abspath(os.path.join(root, f))
            ret.append(cur_file)
    return ret

def find_word_in_docx(word, docx):
    '주어진 docx 파일에서 주어진 word을 찾는다. 찾으면 True를 반환한다.'

    # 문서 로드
    document = Document(docx)

    # 모든 문단을 찾는다.
    for p in document.paragraphs:
        for i in p.runs:
            # 각 run의 text값을 바꾼다.
            s = i.text
            if word in s: return True;

    # 모든 테이블을 찾는다.
    for t in document.tables:
        # 모든 cell에 있는 
        for r in t.rows:
            for c in r.cells:
                # 문단을 찾아서 
                for i in c.paragraphs:
                    for j in i.runs:
                        # 각 run의 text값을 바꾼다.
                        s = j.text
                        if word in s: return True;

    return False


# 테스트 코드
# 현재 디렉토리 아래 "파이썬" 이라는 단어가 들어간 문서를 찾는다.
target_dir = '.'
keywords = '파이썬'

for f in find_docx(target_dir):
    if(find_word_in_docx(keywords, f)):
        print( '- {}'.format(f) );
